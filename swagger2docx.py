
import json

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement

obj = json.load(open('swagger.json'))

document = Document()

defs_dict = obj['definitions']
defs = defs_dict.keys()

path_dict = obj['paths']

def _set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd') # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    if color:
        pass # TODO
    if val:
        pass # TODO
    cell_properties.append(cell_shading)  # finally extend cell props with shading element

def _set_cell_foreground(cell, color):
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)

def _set_cell_foreground_method(cell, method='get'):
    cell.paragraphs[0].runs[0].bold = True
    if method.lower() == 'get':
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('0151A1')
    elif method.lower() == 'post':
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('00682F')
    elif method.lower() == 'put':
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('FC9518')
    elif method.lower() == 'delete':
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('F82020')

def _set_cell_width(cell, width):
    cell.width = Inches(width)

def _set_cell_font_bold(cell):
    cell.paragraphs[0].runs[0].font.bold = True

def _set_paragraph_background(paragraph, fill):
    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)

    # Make sure the paragraph styling element exists
    paragraph.paragraph_format.element.get_or_add_pPr()

    # Append the shading element
    paragraph.paragraph_format.element.pPr.append(shd)

def _get_sample_object_rec(d, is_array = False):
    print(d)
    if not d in defs_dict:
        return None

    sample = dict()

    obj = defs_dict[d]

    if 'type' in obj:
        if obj['type'] == 'object':
            
            if 'properties' in obj:
                props = obj['properties']

                for prop in props.keys():
                    if 'type' in props[prop]:
                        # if array
                        if props[prop]['type'] == 'array':
                            pass
                        elif props[prop]['type'] == 'object':
                            pass
                        else:
                            if 'default' in props[prop]:
                                sample[prop] = props[prop]['default']
                            else: # replace with int, string etc.
                                sample[prop] = '""'
                    elif '$ref' in props[prop]:
                        #print(prop)
                        temp = props[prop]['$ref'].split('/')[-1]
                        sample[prop] = _get_sample_object_rec(temp, False)
        elif obj['type'] == 'array':
            if  'items' in obj:
                if '$ref' in obj['items']:
                    temp = obj['items']['$ref'].split('/')[-1]
                    #print(d, temp, obj['type'])
                    sample = _get_sample_object_rec(temp, True)
    elif '$ref' in obj:
        temp = obj['$ref'].split('/')[-1]
        sample = _get_sample_object_rec(temp, False)
                

    if is_array:
        temp = []
        temp.append(sample)
        sample = temp

    return sample

def _get_sample_object(d, is_array = False):
    if not d in defs_dict:
        return None

    sample = dict()

    obj = defs_dict[d]

    if 'type' in obj:
        if obj['type'] == 'object':
            if 'properties' in obj:
                props = obj['properties']

                for prop in props.keys():
                    if 'type' in props[prop]:
                        # if array
                        if props[prop]['type'] == 'array':
                            # if simple object
                            if 'items' in props[prop]:
                                if '$ref' in props[prop]['items']:
                                    temp = props[prop]['items']['$ref'].split('/')[-1]
                                    sample[prop] = _get_sample_object_rec(temp, True)
                                else:
                                    if 'default' in props[prop]['items']:
                                        sample[prop] = list([props[prop]['items']['default']])
                                    else:
                                        pass # todo: if no default value then default

                        # elif props[prop]['type'] == 'object':
                        #     temp = _get_sample_object_rec(props[prop], False)
                        #     sample[prop] = temp
                        else:
                            if 'default' in props[prop]:
                                sample[prop] = props[prop]['default']
                            else: # replace with int, string etc.
                                sample[prop] = '""'
                    elif '$ref' in props[prop]:
                        temp = props[prop]['$ref'].split('/')[-1]
                        sample[prop] = _get_sample_object_rec(temp, False)

    if is_array:
        temp = []
        temp.append(sample)
        sample = temp

    return json.dumps(sample, indent = 4)

obj_styles = document.styles
obj_charstyle = obj_styles.add_style('CODE', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Consolas'

for path in path_dict:

    print('path:', path)
    methods = path_dict[path].keys()
    for method in methods:
        print(' method:', method)
        resource = path_dict[path][method]
        print('  ', resource['summary'])
        print('    ', resource['description'])

        is_jwt = False

        document.add_heading(resource['summary'], level=2)
        document.add_paragraph(resource['description'])

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        table.columns[0].width = Inches(1.5)

        row_cells = table.rows[0].cells
        row_cells[0].text = 'Path'
        _set_cell_font_bold(row_cells[0])
        _set_cell_background(row_cells[0], 'BDD6EE')
        _set_cell_width(row_cells[0], 1.5)
        row_cells[1].text = path
        row_cells[1].paragraphs[0].style='CODE'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Method'
        _set_cell_font_bold(row_cells[0])        
        _set_cell_width(row_cells[0], 1.5)
        _set_cell_background(row_cells[0], 'BDD6EE')
        row_cells[1].text = method.upper()
        row_cells[1].paragraphs[0].style='CODE'
        _set_cell_foreground_method(row_cells[1], method)

        sample_request = ''

        if 'parameters' in resource or 'security' in resource:
            document.add_heading('Request', level=3)

            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Type'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Description'
            hdr_cells[3].text = 'Schema/Value'
            
            _set_cell_background(hdr_cells[0], 'BDD6EE')
            _set_cell_background(hdr_cells[1], 'BDD6EE')
            _set_cell_background(hdr_cells[2], 'BDD6EE')
            _set_cell_background(hdr_cells[3], 'BDD6EE')

            _set_cell_font_bold(hdr_cells[0])
            _set_cell_font_bold(hdr_cells[1])
            _set_cell_font_bold(hdr_cells[2])
            _set_cell_font_bold(hdr_cells[3])


            # headers
            if 'security' in resource:
                if 'BasicAuth' in resource['security'][0]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'HEADERS'
                    row_cells[1].text = 'Authorization'
                    row_cells[2].text = 'Basic Authentication'
                    row_cells[3].text = 'Basic [userid:password in base64]'

                elif 'JWT' in resource['security'][0]:
                    is_jwt = True
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'HEADERS'
                    row_cells[1].text = 'Authorization'
                    row_cells[2].text = 'JWT Authentication'
                    row_cells[3].text = 'Bearer [token]'

            if 'consumes' in resource:
                for ct in resource['consumes']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'HEADERS'
                    row_cells[1].text = 'Content-Type'
                    row_cells[2].text = 'Content Type of Request'
                    row_cells[3].text =  ct

            if 'parameters' in resource:
                sample_request = ''

                for param in resource['parameters']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = param['in'].upper()
                    row_cells[1].text = param['name']
                    row_cells[2].text = param['description']

                    is_schema = False
                    is_array = False

                    if 'type' in param:
                        row_cells[3].text = param['type']
                        if param['type'] == "integer" and 'format' in param:
                            row_cells[3].text = param['type'] + ' ({})'.format(param['format'])
                            
                    elif 'schema' in param:
                        obj = param['schema']['$ref'].split('/')[-1]
                        row_cells[3].text = '{} [object]'.format(obj)
                        is_schema = True
                        is_array = False

                    if is_schema:
                        sample_request = _get_sample_object(obj, is_array)


        if 'responses' in resource:
            document.add_heading('Response', level=3)

            if 'produces' in resource:
                for ct in resource['produces']:
                    p = document.add_paragraph()
                    p.add_run('Content-Type: ')
                    p.add_run(ct).font.name = 'Consolas'
                    

            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'HTTP Code'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Schema'

            _set_cell_background(hdr_cells[0], 'BDD6EE')
            _set_cell_background(hdr_cells[1], 'BDD6EE')
            _set_cell_background(hdr_cells[2], 'BDD6EE')

            _set_cell_font_bold(hdr_cells[0])
            _set_cell_font_bold(hdr_cells[1])
            _set_cell_font_bold(hdr_cells[2])

            samples = dict()

            responses = resource['responses']
            for code in responses.keys():
                row_cells = table.add_row().cells
                row_cells[0].text = code
                
                row_cells[1].text = responses[code]['description']
                row_cells[2].text = 'No Content'

                is_schema = False
                is_array = False
                if 'schema' in responses[code]:
                    if '$ref' in responses[code]['schema']:
                        obj = responses[code]['schema']['$ref'].split('/')[-1]
                        row_cells[2].text = '{} [object]'.format(obj)
                        is_schema = True
                    elif 'items' in responses[code]['schema']:
                        obj = responses[code]['schema']['items']['$ref'].split('/')[-1]
                        row_cells[2].text = '{} [array]'.format(obj)
                        is_schema = True
                        is_array = True

                if code == '200':
                    _set_cell_foreground(row_cells[0], '00682F')
                    _set_cell_foreground(row_cells[1], '00682F')
                    _set_cell_foreground(row_cells[2], '00682F')
                    _set_cell_font_bold(row_cells[0])
                    _set_cell_font_bold(row_cells[1])
                    _set_cell_font_bold(row_cells[2])

                if is_schema:
                    sample = _get_sample_object(obj, is_array)
                    samples[code] = sample

            # if JWT then add 401 and 403
            if is_jwt:
                row_cells = table.add_row().cells
                row_cells[0].text = '401'
                row_cells[1].text = 'Unauthorized'
                row_cells[2].text = 'No Content'

                row_cells = table.add_row().cells
                row_cells[0].text = '403'
                row_cells[1].text = 'Forbidden'
                row_cells[2].text = 'No Content'

        if sample_request and len(sample_request) > 0:
            document.add_heading('Sample Request [body]', level=3)

            p = document.add_paragraph(sample_request, style='CODE')
            _set_paragraph_background(p, 'E0E0E0')

        if samples and len(samples) > 0:
            document.add_heading('Sample Response [body]', level=3)
            for k, s in samples.items():
                p = document.add_paragraph('HTTP Code: {}'.format(k))
                p = document.add_paragraph(s, style='CODE')
                _set_paragraph_background(p, 'E0E0E0')

document.add_heading('Models', level=2)

c = RGBColor.from_string('274467')

for d in defs:
    print(d)
    s = []
    document.add_heading(d, level=3)
    if 'type' in defs_dict[d]:
        if defs_dict[d]['type'] == 'object':
            props_dict = defs_dict[d]['properties']
            props = props_dict.keys()
            
            for prop in props:
                if 'type' in props_dict[prop]:
                    if props_dict[prop]['type'] == 'array':
                        if 'items' in props_dict[prop]:
                            if 'format' in props_dict[prop]['items']:
                                #s += '{}[] {}\n'.format(props_dict[prop]['items']['format'], prop)
                                s.append([props_dict[prop]['items']['format'], prop])
                            elif 'type' in props_dict[prop]['items']:
                                #s += '{}[] {}\n'.format(props_dict[prop]['items']['type'], prop)
                                s.append([props_dict[prop]['items']['type'], prop])
                            elif '$ref' in props_dict[prop]['items']:
                                temp = props_dict[prop]['items']['$ref'].split('/')[-1]
                                #s += '{}[] {}\n'.format(temp, prop)
                                s.append([temp, prop])
                            else:
                                print('error')
                    else:
                        if 'format' in props_dict[prop]:
                            #s += '{} {}\n'.format(props_dict[prop]['format'], prop)
                            s.append([props_dict[prop]['format'], prop])
                        else:
                            #s += '{} {}\n'.format(props_dict[prop]['type'], prop)
                            s.append([props_dict[prop]['type'], prop])
                elif '$ref' in props_dict[prop]:
                    temp = props_dict[prop]['$ref'].split('/')[-1]
                    print(' ', prop)
                    #s += '{} {}\n'.format(temp, prop)
                    s.append([temp, prop])

            

        elif defs_dict[d]['type'] == 'array':
            pass
    elif '$ref' in defs_dict[d]:
        temp = defs_dict[d]['$ref'].split('/')[-1]
        #s += '{} {}\n'.format(temp, d)
        s.append([temp, d])
    else:
        pass

    
    is_first = True
    if len(s) > 0:
        #p = document.add_paragraph(s.strip(), style='CODE')
        p = document.add_paragraph(style='CODE')
        for row in s:
            if is_first:
                is_first = False
            else:
                p.add_run('\n')
            p.add_run(row[0]).font.color.rgb = c
            p.add_run(' ')
            p.add_run(row[1]).bold = True            
        _set_paragraph_background(p, 'E0E0E0')

    


document.save('swagger.docx')
